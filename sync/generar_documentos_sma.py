"""Genera los DOCX SMA y, si LibreOffice está disponible, sus PDF."""
from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path

from docx import Document


def dni_seguro(value: object) -> str:
    dni = re.sub(r"\D", "", str(value))
    if not 6 <= len(dni) <= 12:
        raise ValueError("DNI/documento inválido")
    return dni


def fecha_larga(value: str) -> str:
    months = ("enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre")
    parsed = date.fromisoformat(value)
    return f"{parsed.day:02d} de {months[parsed.month - 1]} de {parsed.year}"


def set_run(paragraph, index: int, value: object) -> None:
    paragraph.runs[index].text = "" if value is None else str(value)


def set_cell(cell, value: str) -> None:
    paragraph = cell.paragraphs[-1]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def generar_cesion(template: Path, output: Path, data: dict) -> Path:
    doc = Document(template)
    p = doc.paragraphs
    set_run(p[1], 0, f"Lima, {fecha_larga(data['fecha_documento'])}")
    set_run(p[3], 3, data["nombre_completo"])
    set_run(p[3], 10, dni_seguro(data["dni"]))
    set_run(p[3], 16, data["direccion"])
    set_run(p[3], 22, data["distrito"])
    set_run(p[3], 29, data["departamento"])
    set_run(p[4], 2, data["numero_operacion"])
    set_run(p[6], 4, f"{float(data['deuda_total']):.2f}")
    # La fecha legal "23 de junio de 2025" permanece intacta en el párrafo 8.
    destination = output / f"Cesión de Derechos - DNI_{dni_seguro(data['dni'])}.docx"
    doc.save(destination)
    return destination


def generar_acuerdo(template: Path, output: Path, data: dict) -> Path:
    installments = data.get("cuotas") or []
    if not installments:
        raise ValueError("El acuerdo debe contener al menos una cuota")
    agreed = round(float(data["monto_acordado"]), 2)
    if round(sum(float(x["monto"]) for x in installments) + float(data.get("monto_inicial", 0)), 2) != agreed:
        raise ValueError("Las cuotas no suman el monto acordado")
    doc = Document(template)
    p = doc.paragraphs
    set_run(p[9], 5, fecha_larga(data["fecha_documento"]))
    set_run(p[11], 2, data["nombre_completo"])
    set_run(p[12], 3, dni_seguro(data["dni"]))
    set_run(p[13], 2, data.get("telefono") or "-")
    set_run(p[13], 6, data.get("correo") or "-")
    set_run(p[16], 1, " " + str(data.get("producto") or "TARJETA"))
    set_run(p[17], 4, data["numero_operacion"])
    summary = doc.tables[0].rows[1].cells
    set_cell(summary[0], str(data["numero_operacion"]))
    set_cell(summary[1], f"S/ {float(data['deuda_total']):,.2f}")
    set_cell(summary[2], f"S/ {agreed:,.2f}")
    set_cell(summary[3], f"{len(installments):02d}")
    set_cell(summary[4], "CANCELACIÓN TOTAL" if data.get("tipo_acuerdo") == "CANCELACION" else "CONVENIO")
    table = doc.tables[1]
    template_row = deepcopy(table.rows[1]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for number, installment in enumerate(installments, 1):
        new_row = deepcopy(template_row)
        table._tbl.append(new_row)
        cells = table.rows[-1].cells
        set_cell(cells[0], f"{number:02d}")
        parsed = datetime.strptime(installment["fecha"], "%Y-%m-%d")
        set_cell(cells[1], parsed.strftime("%d/%m/%Y"))
        set_cell(cells[2], f"S/ {float(installment['monto']):,.2f}")
    destination = output / f"Acuerdo de pago - DNI_{dni_seguro(data['dni'])}.docx"
    doc.save(destination)
    return destination


def convert_pdf(docx: Path, output: Path, soffice: str) -> Path:
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output), str(docx)], check=True, capture_output=True, text=True, timeout=120)
    generated = output / f"{docx.stem}.pdf"
    if not generated.is_file() or generated.stat().st_size == 0:
        raise RuntimeError(f"LibreOffice no generó {generated.name}")
    return generated


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, type=Path)
    parser.add_argument("--templates", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--soffice", default=shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.com")
    parser.add_argument("--docx-only", action="store_true")
    args = parser.parse_args()
    data = json.loads(args.data.read_text(encoding="utf-8"))
    args.output.mkdir(parents=True, exist_ok=True)
    files = [
        generar_acuerdo(args.templates / "Acuerdo_de_Pago.docx", args.output, data),
        generar_cesion(args.templates / "Cesion_Derechos.docx", args.output, data),
    ]
    if not args.docx_only:
        files += [convert_pdf(item, args.output, args.soffice) for item in files]
    print(json.dumps({"ok": True, "files": [str(x) for x in files]}, ensure_ascii=False))


if __name__ == "__main__":
    main()

