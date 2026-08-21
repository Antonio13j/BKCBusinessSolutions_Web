import json
import tempfile
import unittest
from pathlib import Path
from docx import Document
from sync.generar_documentos_sma import generar_acuerdo, generar_cesion

ROOT=Path(__file__).resolve().parent
TEMPLATES=Path(r"C:\VS_Code\Cesion_Derechos\Plantillas")

class DocumentosSmaTest(unittest.TestCase):
    def setUp(self):
        self.data=json.loads((ROOT/'sample_sma_demo.json').read_text(encoding='utf-8'))

    def test_cinco_cuotas_y_nombre(self):
        with tempfile.TemporaryDirectory() as folder:
            result=generar_acuerdo(TEMPLATES/'Acuerdo_de_Pago.docx',Path(folder),self.data)
            self.assertEqual(result.name,'Acuerdo de pago - DNI_00000000.docx')
            doc=Document(result)
            self.assertEqual(len(doc.tables[1].rows),6)
            self.assertEqual([r.cells[0].text for r in doc.tables[1].rows[1:]],['01','02','03','04','05'])
            self.assertEqual(doc.tables[1].rows[-1].cells[2].text,'S/ 87.00')

    def test_cesion_conserva_fecha_legal(self):
        with tempfile.TemporaryDirectory() as folder:
            result=generar_cesion(TEMPLATES/'Cesion_Derechos.docx',Path(folder),self.data)
            self.assertEqual(result.name,'Cesión de Derechos - DNI_00000000.docx')
            doc=Document(result)
            self.assertIn('23 de junio de 2025',doc.paragraphs[8].text)
            self.assertIn('00000000',doc.paragraphs[3].text)

if __name__=='__main__':unittest.main()
